"""
学习资料模块
管理上传和关联学习资料
"""

import streamlit as st
from utils import render_html
import os
from datetime import datetime

from db import (
    get_learning_materials, create_learning_material, get_material_by_id,
    delete_learning_material, get_skill_list, update_material_position,
    get_materials_stats, update_material_content
)


def extract_file_content(file_path: str, file_type: str) -> str:
    """
    根据文件类型提取文本内容
    
    Args:
        file_path: 文件完整路径
        file_type: 文件类型（扩展名）
    
    Returns:
        str: 提取的文本内容，失败则返回空字符串
    """
    file_type = file_type.lower()
    
    try:
        if file_type == 'pdf':
            return extract_pdf_text(file_path)
        elif file_type in ['txt', 'md']:
            return extract_text_file(file_path)
        elif file_type in ['docx', 'doc']:
            return extract_docx_text(file_path)
        else:
            return "暂不支持提取此类文件的文本内容"
    except Exception as e:
        return f"文件内容提取失败: {str(e)}"


def extract_pdf_text(file_path: str) -> str:
    """提取PDF文件文本，依次尝试多种方法"""
    text_parts = []

    # 方法1: pdfplumber（最强）
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        if text_parts:
            return '\n'.join(text_parts)
    except ImportError:
        pass
    except Exception:
        pass

    # 方法2: PyPDF2
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        if text_parts:
            return '\n'.join(text_parts)
    except ImportError:
        pass
    except Exception:
        pass

    # 方法3: pymupdf (fitz)
    try:
        import fitz
        doc = fitz.open(file_path)
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text_parts.append(page_text)
        doc.close()
        if text_parts:
            return '\n'.join(text_parts)
    except ImportError:
        pass
    except Exception:
        pass

    if not text_parts:
        return "该PDF无法提取文本（可能是扫描版图片PDF，需要OCR才能识别文字）。建议：1)安装pdfplumber库 2)或手动将关键内容复制粘贴到笔记中"
    return '\n'.join(text_parts)


def extract_text_file(file_path: str) -> str:
    """提取文本文件内容"""
    try:
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        return "无法解码文本文件"
    except Exception as e:
        return f"文本文件读取失败: {str(e)}"


def extract_docx_text(file_path: str) -> str:
    """提取DOCX文件文本"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n'.join(paragraphs)
    except ImportError:
        return "需要安装 python-docx 库来提取 DOCX 内容"
    except Exception as e:
        return f"DOCX 提取失败: {str(e)}"


def get_theme_colors() -> dict:
    """获取当前主题颜色"""
    is_dark = st.session_state.get('theme', 'dark') == 'dark'
    if is_dark:
        return {
            'bg_secondary': '#1b263b',
            'bg_tertiary': '#334155',
            'text_primary': '#e0e1dd',
            'text_secondary': '#a0aec0',
            'text_muted': '#64748b',
            'accent': '#60a5fa',
            'success': '#4ade80',
            'warning': '#facc15',
            'danger': '#ef4444',
            'purple': '#a855f7',
            'card_border': '#415a77',
        }
    else:
        return {
            'bg_secondary': '#ffffff',
            'bg_tertiary': '#e2e8f0',
            'text_primary': '#1e293b',
            'text_secondary': '#475569',
            'text_muted': '#94a3b8',
            'accent': '#3b82f6',
            'success': '#22c55e',
            'warning': '#eab308',
            'danger': '#ef4444',
            'purple': '#9333ea',
            'card_border': '#cbd5e1',
        }


def render_materials_page():
    """
    渲染学习资料页面
    """
    colors = get_theme_colors()
    is_guest = st.session_state.get('guest_mode', False)
    
    # 页面标题
    render_html(f"""
    <div class="welcome-section" style="padding: 1.5rem;">
        <h1 style="margin: 0; color: {colors['text_primary']}; font-size: 1.8rem;">📚 学习资料</h1>
        <p style="margin: 0.5rem 0 0 0; color: {colors['text_secondary']};">管理你的学习文件和资料</p>
    </div>
    """)
    
    # 获取统计数据
    stats = get_materials_stats()
    
    # 统计信息
    total_size = stats['total_size'] or 0
    if total_size > 1024 * 1024:
        size_str = f"{total_size / (1024 * 1024):.2f} MB"
    elif total_size > 1024:
        size_str = f"{total_size / 1024:.2f} KB"
    else:
        size_str = f"{total_size} B"
    
    stats_html = f"""
    <div style="display: grid; grid-template-columns: repeat(2, 1fr); gap: 1rem; margin-bottom: 1.5rem;">
        <div class="stat-card">
            <div class="stat-icon">📦</div>
            <div class="stat-value" style="color: {colors['accent']};">{stats['total']}</div>
            <div class="stat-label">资料总数</div>
        </div>
        <div class="stat-card">
            <div class="stat-icon">💾</div>
            <div class="stat-value" style="color: {colors['purple']};">{size_str}</div>
            <div class="stat-label">总大小</div>
        </div>
    </div>
    """
    render_html(stats_html)
    
    # 上传资料区域
    if not is_guest:
        render_upload_section()
    
    # 筛选
    st.markdown("### 📋 资料列表")
    
    skill_filter = st.selectbox(
        "按技能筛选",
        options=["全部", "未关联"] + [s['skill_name'] for s in get_skill_list()],
        key="material_skill_filter"
    )
    
    # 获取资料列表
    skill_name = None if skill_filter == "全部" else skill_filter
    if skill_filter == "未关联":
        materials = [m for m in get_learning_materials() if not m.get('skill_name')]
    else:
        materials = get_learning_materials(skill_name=skill_name)
    
    # 自动重新提取空内容的PDF
    if not is_guest:
        materials_all = get_learning_materials()
        for m in materials_all:
            ct = m.get('content_text', '')
            needs = (not ct or "无法" in ct or "失败" in ct or "不支持" in ct or "需要安装" in ct)
            if needs and m.get('file_type', '').lower() == 'pdf':
                fp = m.get('file_path', '')
                if fp and os.path.exists(fp):
                    try:
                        extracted = extract_file_content(fp, 'pdf')
                        if extracted and len(extracted) > 50:
                            from db import update_material_content
                            update_material_content(m['id'], extracted)
                    except:
                        pass

    if not materials:
        st.info("📭 还没有上传资料！")
    else:
        render_materials_list(materials, is_guest)


def render_upload_section():
    """
    渲染上传资料区域
    """
    colors = get_theme_colors()
    
    st.markdown("### ⬆️ 上传新资料")
    
    # 获取技能列表用于关联
    skills = get_skill_list()
    skill_options = ["未关联"] + [s['skill_name'] for s in skills]
    
    with st.expander("点击上传文件", expanded=False):
        with st.form(key="upload_material_form"):
            uploaded_file = st.file_uploader(
                "选择文件",
                type=['pdf', 'doc', 'docx', 'txt', 'md', 'xlsx', 'xls', 'pptx', 'ppt', 'zip', 'rar'],
                key="file_uploader"
            )
            
            skill_name = st.selectbox(
                "关联技能",
                options=skill_options,
                key="upload_skill_select"
            )
            
            description = st.text_area("文件描述", key="upload_desc", height=60)
            
            submitted = st.form_submit_button("📤 上传", type="primary", use_container_width=True)
            
            if submitted:
                if uploaded_file is None:
                    st.error("请选择要上传的文件")
                else:
                    # 保存文件
                    skill_for_path = skill_name if skill_name != "未关联" else "未关联"
                    upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads", skill_for_path)
                    os.makedirs(upload_dir, exist_ok=True)
                    
                    file_path = os.path.join(upload_dir, uploaded_file.name)
                    
                    # 如果文件已存在，添加时间戳
                    if os.path.exists(file_path):
                        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                        name, ext = os.path.splitext(uploaded_file.name)
                        file_path = os.path.join(upload_dir, f"{name}_{timestamp}{ext}")
                    
                    # 写入文件
                    with open(file_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())
                    
                    # 获取文件类型
                    file_type = uploaded_file.name.split('.')[-1].lower() if '.' in uploaded_file.name else 'unknown'
                    
                    # 提取文件内容
                    with st.spinner("正在提取文件内容..."):
                        content_text = extract_file_content(file_path, file_type)
                    
                    # 保存到数据库
                    skill_to_save = "" if skill_name == "未关联" else skill_name
                    material_id = create_learning_material(
                        skill_name=skill_to_save,
                        file_name=uploaded_file.name,
                        file_path=file_path,
                        file_type=file_type,
                        file_size=uploaded_file.size,
                        description=description,
                        content_text=content_text
                    )
                    
                    # 如果提取的内容是错误信息，更新到数据库
                    if "失败" in content_text or "无法" in content_text or "不支持" in content_text or "需要安装" in content_text:
                        update_material_content(material_id, content_text)
                    
                    st.success(f"文件 '{uploaded_file.name}' 上传成功！")
                    st.rerun()


def render_materials_list(materials: list, is_guest: bool):
    """
    渲染资料列表
    
    Args:
        materials: 资料列表
        is_guest: 是否访客模式
    """
    for material in materials:
        render_material_card(material, is_guest)


def render_material_card(material: dict, is_guest: bool):
    """
    渲染单个资料卡片
    
    Args:
        material: 资料数据
        is_guest: 是否访客模式
    """
    colors = get_theme_colors()
    
    # 文件类型图标
    file_icons = {
        'pdf': '📄',
        'doc': '📝', 'docx': '📝',
        'txt': '📃', 'md': '📃',
        'xlsx': '📊', 'xls': '📊',
        'pptx': '📽️', 'ppt': '📽️',
        'zip': '📦', 'rar': '📦',
        'default': '📁'
    }
    icon = file_icons.get(material.get('file_type', '').lower(), file_icons['default'])
    
    # 格式化文件大小
    file_size = material.get('file_size', 0)
    if file_size > 1024 * 1024:
        size_str = f"{file_size / (1024 * 1024):.2f} MB"
    elif file_size > 1024:
        size_str = f"{file_size / 1024:.2f} KB"
    else:
        size_str = f"{file_size} B"
    
    with st.container():
        card_html = f"""
        <div class="content-card">
            <div style="display: flex; align-items: center;">
                <span style="font-size: 2.5rem; margin-right: 1rem;">{icon}</span>
                <div style="flex: 1;">
                    <h3 style="margin: 0; color: {colors['text_primary']};">{material['file_name']}</h3>
                    <div style="display: flex; gap: 1rem; margin-top: 0.3rem; color: {colors['text_secondary']}; font-size: 0.85rem;">
                        <span>💾 {size_str}</span>
                        <span>🗂️ {material.get('file_type', 'unknown').upper()}</span>
                        <span>📅 {material.get('upload_time', '')[:10]}</span>
                    </div>
        """
        
        if material.get('skill_name'):
            card_html += f"""
                    <div style="margin-top: 0.5rem;">
                        <span style="background: {colors['accent']}; padding: 0.25rem 0.6rem; border-radius: 6px; 
                                    font-size: 0.8rem; color: white;">
                            🎯 {material['skill_name']}
                        </span>
                    </div>
            """
        
        if material.get('description'):
            card_html += f"""
                    <p style="color: {colors['text_muted']}; margin: 0.5rem 0 0 0; font-size: 0.85rem;">
                        💬 {material['description']}
                    </p>
            """
        
        card_html += """
                </div>
            </div>
        </div>
        """
        render_html(card_html)
        
        # 操作按钮
        btn_col1, btn_col2, btn_col3, btn_col4 = st.columns(4)

        with btn_col1:
            preview_key = f"_preview_mat_{material['id']}"
            if st.session_state.get(preview_key, False):
                if st.button("✕ 关闭", key=f"btn_close_{material['id']}", type="primary", use_container_width=True):
                    st.session_state[preview_key] = False
                    st.rerun()
            else:
                if st.button("👁️ 预览", key=f"btn_preview_{material['id']}", type="primary", use_container_width=True):
                    st.session_state[preview_key] = True
                    st.rerun()

        with btn_col2:
            file_path = material.get('file_path')
            if file_path and os.path.exists(file_path):
                try:
                    with open(file_path, 'rb') as f:
                        file_data = f.read()
                    st.download_button(
                        label="📥 下载",
                        data=file_data,
                        file_name=material['file_name'],
                        mime='application/octet-stream',
                        key=f"download_btn_{material['id']}",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"无法读取文件: {str(e)}")
            else:
                st.button("📥 下载", disabled=True, use_container_width=True, key=f"dl_disabled_{material['id']}")

        with btn_col3:
            if st.button("🔗 继续", key=f"continue_material_{material['id']}", use_container_width=True):
                skill_name = material.get('skill_name')
                if skill_name:
                    skills = get_skill_list(search=skill_name)
                    if skills:
                        st.session_state['view_item_id'] = skills[0]['id']
                        st.session_state['view_item_type'] = 'skill'
                        st.session_state['current_page'] = 'growth'
                        st.session_state['open_material_id'] = material['id']
                        st.rerun()
                else:
                    st.info("这个资料没有关联技能")

        if not is_guest:
            with btn_col4:
                # 如果没有文本内容或提取失败，显示"重新提取"按钮
                content_text = material.get('content_text', '')
                needs_extract = (not content_text or 
                    "无法" in (content_text or "") or 
                    "失败" in (content_text or "") or 
                    "不支持" in (content_text or "") or
                    "需要安装" in (content_text or "") or
                    "需要OCR" in (content_text or ""))
                if needs_extract:
                    if st.button("📝 提取", key=f"extract_material_{material['id']}", use_container_width=True):
                        file_path_tmp = material.get('file_path')
                        file_type_tmp = material.get('file_type', '')
                        if file_path_tmp and os.path.exists(file_path_tmp):
                            with st.spinner("正在提取文本..."):
                                extracted = extract_file_content(file_path_tmp, file_type_tmp)
                            from db import update_material_content
                            update_material_content(material['id'], extracted)
                            st.success("文本提取完成！")
                            st.rerun()
                        else:
                            st.error("文件不存在")
                if st.button("🗑️ 删除", key=f"delete_material_{material['id']}", use_container_width=True):
                    file_path = material.get('file_path')
                    if file_path and os.path.exists(file_path):
                        try:
                            os.remove(file_path)
                        except:
                            pass
                    delete_learning_material(material['id'])
                    st.success("资料已删除")
                    st.rerun()

        # 预览区域
        preview_key = f"_preview_mat_{material['id']}"
        if st.session_state.get(preview_key, False):
            file_path = material.get('file_path')
            file_type = material.get('file_type', '').lower()

            st.markdown("---")

            if file_path and os.path.exists(file_path):
                try:
                    if file_type == 'pdf':
                        # PDF展示：提取文本 + 下载提示
                        from db import get_material_content
                        text_content = get_material_content(material['id'])
                        st.markdown(f"**📄 {material['file_name']}**")
                        
                        tab_text, tab_download = st.tabs(["📝 文本内容", "📥 下载原文"])
                        with tab_text:
                            if text_content and len(text_content) > 50:
                                st.info("💡 以下为自动提取的文本，数学公式可能显示为乱码，完整公式请下载原文查看")
                                st.text_area("文本内容", text_content, height=500, key=f"pdf_text_{material['id']}")
                            else:
                                st.warning("未提取到文本内容，可能是扫描版PDF")
                        with tab_download:
                            file_path_dl = material.get('file_path')
                            if file_path_dl and os.path.exists(file_path_dl):
                                with open(file_path_dl, 'rb') as f:
                                    file_data = f.read()
                                st.download_button(
                                    label="📥 下载PDF原文",
                                    data=file_data,
                                    file_name=material['file_name'],
                                    mime='application/pdf',
                                    key=f"preview_download_{material['id']}",
                                    use_container_width=True
                                )
                                st.markdown("**提示**：下载后在本地PDF阅读器中查看，公式和排版都完整")
                    elif file_type in ('txt', 'md', 'py', 'csv', 'json', 'js', 'html', 'css', 'log'):
                        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                            text_content = f.read()
                        st.text_area("文件内容", text_content, height=400, key=f"preview_text_{material['id']}")
                    elif file_type in ('docx', 'doc'):
                        try:
                            from docx import Document
                            doc = Document(file_path)
                            full_text = '\n'.join([para.text for para in doc.paragraphs])
                            st.text_area("文档内容", full_text, height=400, key=f"preview_docx_{material['id']}")
                        except ImportError:
                            st.warning("需要安装 python-docx 才能预览Word文档")
                    else:
                        st.info(f"暂不支持在线预览 {file_type.upper()} 格式，请下载后查看")
                except Exception as e:
                    st.error(f"预览失败: {str(e)}")
            else:
                st.error("文件不存在，可能已被移动或删除")
            st.markdown("---")


def format_file_size(size_bytes: int) -> str:
    """
    格式化文件大小
    
    Args:
        size_bytes: 字节数
    
    Returns:
        str: 格式化后的大小字符串
    """
    if size_bytes > 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.2f} MB"
    elif size_bytes > 1024:
        return f"{size_bytes / 1024:.2f} KB"
    else:
        return f"{size_bytes} B"
